import os
import uuid
import datetime
import copy
import re
import unicodedata
from docx import Document
from docx.shared import RGBColor
from docx.table import _Row
from docx.oxml import OxmlElement
from django.conf import settings

def format_non_heading_text(doc):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    font_name = "Arial"
    font_size = Pt(11)
        
    # Apply to all non-heading paragraphs in body
    for p in doc.paragraphs:
        style_lower = p.style.name.lower()
        if not (style_lower.startswith("heading") or style_lower in ("title", "subtitle")):
            p_text_norm = normalize_text(p.text)
            is_vco_disclaimer = "os registros do vco" in p_text_norm or ("canal de etica" in p_text_norm and "14.457" in p_text_norm)
            for r in p.runs:
                r.font.name = font_name
                r.font.size = font_size
                if is_vco_disclaimer:
                    r.italic = True
                    
    # Apply to all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    style_lower = p.style.name.lower()
                    if not (style_lower.startswith("heading") or style_lower in ("title", "subtitle")):
                        p_text_norm = normalize_text(p.text)
                        is_vco_disclaimer = "os registros do vco" in p_text_norm or ("canal de etica" in p_text_norm and "14.457" in p_text_norm)
                        for r in p.runs:
                            r.font.name = font_name
                            r.font.size = font_size
                            if is_vco_disclaimer:
                                r.italic = True
                        if p.alignment not in (1, 2):
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def map_copsoq_to_report_class(copsoq_class, dim_code=None):
    if dim_code in ("DEM", "EST", "CTV"):
        mapping = {
            "Favoravel": "Baixo Risco",
            "Atencao": "Baixo Risco",
            "Critico": "Médio Risco",
            "Grave": "Alto Risco",
        }
    else:
        mapping = {
            "Favoravel": "Baixo Risco",
            "Atencao": "Baixo Risco",
            "Critico": "Médio Risco",
            "Grave": "Alto Risco",
        }
    return mapping.get(copsoq_class, "Médio Risco")


def normalize_text(text: str) -> str:
    if not text:
        return ""
    nfkd_form = unicodedata.normalize('NFKD', text)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)]).lower().strip()

def normalize_activity_q6(val: str) -> str:
    if not val:
        return ""
    val_norm = normalize_text(val)
    if "administrativo" in val_norm or "escritorio" in val_norm or "escritório" in val_norm:
        return "Ambiente administrativo (escritório)"
    if "industrial" in val_norm or "producao" in val_norm or "fabrica" in val_norm or "fábrica" in val_norm or "produção" in val_norm:
        return "Ambiente industrial (produção / fábrica)"
    if "logistico" in val_norm or "logístico" in val_norm or "estoque" in val_norm or "almoxarifado" in val_norm:
        return "Ambiente logístico (estoque / almoxarifado)"
    if "hospitalar" in val_norm or "assistencial" in val_norm or "saude" in val_norm or "saúde" in val_norm:
        return "Ambiente hospitalar / assistencial"
    if "externo" in val_norm or "rua" in val_norm or "campo" in val_norm or "deslocamento" in val_norm:
        return "Ambiente externo (rua, campo, deslocamento)"
    if "publico" in val_norm or "público" in val_norm or "atendimento" in val_norm:
        return "Ambiente com atendimento direto ao público"
    return val

def normalize_headcount_key(key: str) -> str:
    if not key:
        return ""
    if " | " in key:
        ghe, act = key.split(" | ", 1)
    elif "|" in key:
        ghe, act = key.split("|", 1)
    else:
        ghe, act = key, ""
    
    ghe_norm = normalize_text(ghe).replace(" ", "")
    act_norm = normalize_text(act).replace(" ", "")
    return f"{ghe_norm}|{act_norm}"

def delete_paragraph(p):
    try:
        p.text = ""
        p_elem = p._element
        p_elem.getparent().remove(p_elem)
    except Exception:
        pass

def copy_paragraph_format(src_para, dst_para):
    try:
        src_pr = src_para._element.get_or_add_pPr()
        dst_pr = dst_para._element.get_or_add_pPr()
        for child in list(dst_pr):
            dst_pr.remove(child)
        for child in src_pr:
            dst_pr.append(copy.deepcopy(child))
    except Exception:
        pass

def add_row_from_xml(table, template_tr):
    new_tr = copy.deepcopy(template_tr)
    table._tbl.append(new_tr)
    new_row = _Row(new_tr, table)
    for cell in new_row.cells:
        cell.text = ""
    return new_row

def replace_regex_everywhere(doc, pattern_str, replacement_text):
    regex = re.compile(pattern_str, re.IGNORECASE)
    
    def process_paragraph(p):
        if not regex.search(p.text):
            return
        # Try run-by-run first to keep formatting
        replaced = False
        for r in p.runs:
            if regex.search(r.text):
                r.text = regex.sub(replacement_text, r.text)
                replaced = True
        # Fallback to paragraph-level replacement
        if not replaced:
            p.text = regex.sub(replacement_text, p.text)

    # 1. Paragraphs
    for p in doc.paragraphs:
        process_paragraph(p)
            
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
                        
    # 3. Sections (Headers/Footers)
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for p in header.paragraphs:
                    process_paragraph(p)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                process_paragraph(p)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer:
                for p in footer.paragraphs:
                    process_paragraph(p)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                process_paragraph(p)

def style_class_cell(cell, classification):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(classification)
    run.bold = True
    
    if "Alto" in classification or "Grave" in classification:
        run.font.color.rgb = RGBColor(192, 0, 0) # Dark Red
    elif "Médio" in classification or "Critico" in classification:
        run.font.color.rgb = RGBColor(237, 125, 49) # Orange
    elif "Fator Protetor" in classification or "Favoravel" in classification or "Baixo" in classification:
        run.font.color.rgb = RGBColor(112, 173, 71) # Green
    else:
        run.font.color.rgb = RGBColor(0, 0, 0) # Black

def populate_section_bullets(doc, heading_text_norm, list_items):
    from docx.text.paragraph import Paragraph
    
    heading_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(heading_text_norm):
            heading_idx = idx
            break
            
    if heading_idx == -1:
        return
        
    example_paras = []
    # Collect existing example paragraphs
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(("6.1", "6.2", "6.3", "6.4", "7.", "8.", "9.", "10.", "11.", "12.", "13.")) or p.style.name.startswith("Heading"):
            break
        if p.text.strip():
            example_paras.append(p)
            
    for i, item in enumerate(list_items):
        if i < len(example_paras):
            example_paras[i].text = item
        else:
            prev_p = example_paras[-1] if example_paras else doc.paragraphs[heading_idx]
            new_p_xml = OxmlElement('w:p')
            prev_p._element.addnext(new_p_xml)
            new_para = Paragraph(new_p_xml, prev_p._parent)
            new_para.text = item
            if example_paras:
                copy_paragraph_format(example_paras[0], new_para)
            else:
                try:
                    new_para.style = "List Bullet"
                except KeyError:
                    try:
                        new_para.style = "ListBullet"
                    except KeyError:
                        try:
                            new_para.style = "Parágrafo com lista"
                        except KeyError:
                            pass
            example_paras.append(new_para)
            
    if len(example_paras) > len(list_items):
        for p in example_paras[len(list_items):]:
            delete_paragraph(p)

def replace_heading_example(doc, heading_text_norm, replacement_text):
    heading_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(heading_text_norm):
            heading_idx = idx
            break
            
    if heading_idx == -1:
        return
        
    to_remove = []
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(("6.1", "6.2", "6.3", "6.4", "7.", "8.", "9.", "10.", "11.", "12.", "13.")) or p.style.name.startswith("Heading"):
            break
        to_remove.append(p)
        
    from docx.text.paragraph import Paragraph
    current_p = doc.paragraphs[heading_idx]
    
    if replacement_text:
        new_p_xml = OxmlElement('w:p')
        current_p._element.addnext(new_p_xml)
        new_para = Paragraph(new_p_xml, current_p._parent)
        new_para.text = replacement_text
        if to_remove:
            copy_paragraph_format(to_remove[0], new_para)
            
    for p in to_remove:
        delete_paragraph(p)

def replace_heading_with_multiple_paragraphs(doc, heading_text_norm, paragraphs_text_list):
    heading_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(heading_text_norm):
            heading_idx = idx
            break
            
    if heading_idx == -1:
        return
        
    to_remove = []
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(("6.1", "6.2", "6.3", "6.4", "7.", "8.", "9.", "10.", "11.", "12.", "13.")) or p.style.name.startswith("Heading"):
            break
        to_remove.append(p)
        
    from docx.text.paragraph import Paragraph
    current_p = doc.paragraphs[heading_idx]
    
    for text in paragraphs_text_list:
        new_p_xml = OxmlElement('w:p')
        current_p._element.addnext(new_p_xml)
        new_para = Paragraph(new_p_xml, current_p._parent)
        new_para.text = text
        if to_remove:
            copy_paragraph_format(to_remove[0], new_para)
        current_p = new_para
        
    for p in to_remove:
        delete_paragraph(p)

def map_matrix_to_report_class(matrix_class, dim_code=None):
    if dim_code in ("DEM", "EST", "CTV"):
        mapping = {
            "Baixo": "Baixo Risco",
            "Moderado": "Baixo Risco",
            "Alto": "Médio Risco",
            "Critico": "Alto Risco",
        }
    else:
        mapping = {
            "Baixo": "Baixo Risco",
            "Moderado": "Baixo Risco",
            "Alto": "Médio Risco",
            "Critico": "Alto Risco",
        }
    return mapping.get(matrix_class, "Médio Risco")


def analyze_dimension_criticality(processed_survey, dimension_code, metodologia):
    item_texts = {}
    q_code = None
    for q_id, q_cfg in metodologia.perguntas.items():
        if q_cfg.get("codigo_dimensao") == dimension_code:
            q_code = q_id
            for item in q_cfg.get("itens", []):
                item_texts[item["codigo"]] = item["texto"]
            break
            
    from collections import Counter
    from psicossocial.workforce import canonical_ghe
    
    item_critical_counts = Counter()
    ghe_critical_counts = Counter()
    ghe_totals = Counter()
    
    for r in processed_survey.respondents:
        # Resolve GHE
        ghe = canonical_ghe(r.characterization.get("Q1", ""))
        ghe_totals[ghe] += 1
        
        dim_score = r.dimensions.get(dimension_code)
        if dim_score:
            for item in dim_score.items:
                if item.risk_score >= 4.0:
                    item_critical_counts[item.item_code] += 1
                    ghe_critical_counts[ghe] += 1
                    
    # Find most critical item
    most_critical_item_code = None
    max_item_count = -1
    for item_code, count in item_critical_counts.items():
        if count > max_item_count:
            max_item_count = count
            most_critical_item_code = item_code
            
    item_desc = ""
    if most_critical_item_code:
        item_desc = item_texts.get(most_critical_item_code, "")
        
    if not item_desc:
        if item_texts:
            item_desc = list(item_texts.values())[0]
        else:
            item_desc = "percepções gerais da dimensão"
            
    # Find GHE with highest prevalence of critical responses
    most_critical_ghe = "Não informado"
    max_ghe_pct = -1.0
    for ghe, crit_count in ghe_critical_counts.items():
        total = ghe_totals[ghe]
        if total > 0:
            pct = crit_count / total
            if pct > max_ghe_pct:
                max_ghe_pct = pct
                most_critical_ghe = ghe
                
    if max_ghe_pct <= 0:
        if ghe_totals:
            most_critical_ghe = ghe_totals.most_common(1)[0][0]
            
    # Format description
    item_desc = item_desc.strip()
    if item_desc:
        if item_desc[0].isupper() and not item_desc.startswith("Q"):
            item_desc = item_desc[0].lower() + item_desc[1:]
        if item_desc.endswith("?"):
            item_desc = item_desc[:-1]
            
    ghe_desc = most_critical_ghe
    if ghe_desc != "Não informado" and not ghe_desc.upper().startswith("GHE"):
        ghe_desc = f"GHE {ghe_desc}"
        
    return item_desc, ghe_desc

def analyze_vco_criticality(processed_survey):
    from collections import Counter
    from psicossocial.workforce import canonical_ghe
    
    item_texts = {
        "Q16_VCO_1": "tratamento desrespeitoso de superiores",
        "Q16_VCO_2": "ameaça verbal ou intimidação",
        "Q16_VCO_3": "assédio moral",
        "Q16_VCO_4": "assédio sexual",
        "Q16_VCO_5": "violência física"
    }
    
    item_critical_counts = Counter()
    ghe_critical_counts = Counter()
    ghe_totals = Counter()
    
    for r in processed_survey.respondents:
        ghe = canonical_ghe(r.characterization.get("Q1", ""))
        ghe_totals[ghe] += 1
        
        for item_code in ("Q16_VCO_1", "Q16_VCO_2", "Q16_VCO_3", "Q16_VCO_4", "Q16_VCO_5"):
            val = r.vco.get(item_code, "").lower().strip()
            if val == "sim":
                item_critical_counts[item_code] += 1
                ghe_critical_counts[ghe] += 1
                
    # Collect all items that were marked "sim" (count > 0), sorted by frequency (most frequent first)
    critical_items = [item_code for item_code, count in item_critical_counts.items() if count > 0]
    critical_items.sort(key=lambda x: item_critical_counts[x], reverse=True)
    
    if critical_items:
        descs = [item_texts[item_code] for item_code in critical_items]
        if len(descs) == 1:
            item_desc = descs[0]
        elif len(descs) == 2:
            item_desc = f"{descs[0]} e {descs[1]}"
        else:
            item_desc = ", ".join(descs[:-1]) + f" e {descs[-1]}"
    else:
        item_desc = "comportamento ofensivo não especificado"
        
    most_critical_ghe = "Não informado"
    max_ghe_pct = -1.0
    for ghe, crit_count in ghe_critical_counts.items():
        total = ghe_totals[ghe]
        if total > 0:
            pct = crit_count / total
            if pct > max_ghe_pct:
                max_ghe_pct = pct
                most_critical_ghe = ghe
                
    if max_ghe_pct <= 0:
        if ghe_totals:
            most_critical_ghe = ghe_totals.most_common(1)[0][0]
            
    ghe_desc = most_critical_ghe
    if ghe_desc != "Não informado" and not ghe_desc.upper().startswith("GHE"):
        ghe_desc = f"GHE {ghe_desc}"
        
    return item_desc, ghe_desc

def append_free_text_to_section_6_3(doc, text):
    if not text:
        return
        
    # Find heading 6.3
    heading_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith("6.3"):
            heading_idx = idx
            break
            
    if heading_idx == -1:
        return
        
    last_p = None
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        p_text_norm = normalize_text(p.text)
        if p_text_norm.startswith(("6.1", "6.2", "6.3", "6.4", "7.", "8.", "9.", "10.", "11.", "12.", "13.")) or p.style.name.startswith("Heading"):
            break
        last_p = p
        
    if last_p:
        new_p_xml = OxmlElement('w:p')
        last_p._element.addnext(new_p_xml)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p_xml, last_p._parent)
        new_para.text = ""
        r1 = new_para.add_run("\nAções corporativas e justificativas de mitigação relatadas:\n")
        r1.bold = True
        r2 = new_para.add_run(text)
        r2.italic = True

def generate_word_report(projeto, processed_survey, metodologia, representatividade_summary, headcounts=None):
    if headcounts is None:
        headcounts = {}
    template_path = os.path.join(settings.BASE_DIR, "psicossocial", "01_entrada_usuario", "Relatorio_Psicossocial_NR01 1.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template Word não encontrado em: {template_path}")
        
    doc = Document(template_path)
    
    # 1. Gather Dimension Stats
    dim_stats = processed_survey.summary(metodologia)["dimensoes"]
    
    from psicossocial.processing import dimension_exposure_summary
    
    def get_matrix_classification(dim_code):
        try:
            exposure = dimension_exposure_summary(processed_survey.respondents, dim_code)
            return exposure.matrix_risk_classification
        except Exception as e:
            print(f"Error calculating matrix exposure for {dim_code}: {e}")
            return "Baixo"
            
    DEM_risk = dim_stats.get("DEM", {}).get("risk_score_medio", 1.0)
    DEM_score = (DEM_risk - 1.0) / 4.0 * 100
    DEM_class = map_matrix_to_report_class(get_matrix_classification("DEM"), "DEM")
    
    EST_risk = dim_stats.get("EST", {}).get("risk_score_medio", 1.0)
    EST_score = (EST_risk - 1.0) / 4.0 * 100
    EST_class = map_matrix_to_report_class(get_matrix_classification("EST"), "EST")
    
    CTV_risk = dim_stats.get("CTV", {}).get("risk_score_medio", 1.0)
    CTV_score = (CTV_risk - 1.0) / 4.0 * 100
    CTV_class = map_matrix_to_report_class(get_matrix_classification("CTV"), "CTV")
    
    REL_risk = dim_stats.get("REL", {}).get("risk_score_medio", 1.0)
    REL_score = (5.0 - REL_risk) / 4.0 * 100
    REL_class = map_matrix_to_report_class(get_matrix_classification("REL"), "REL")
    
    ORG_risk = dim_stats.get("ORG", {}).get("risk_score_medio", 1.0)
    ORG_score = (5.0 - ORG_risk) / 4.0 * 100
    ORG_class = map_matrix_to_report_class(get_matrix_classification("ORG"), "ORG")
    
    ITI_risk = dim_stats.get("ITI", {}).get("risk_score_medio", 1.0)
    ITI_score = (5.0 - ITI_risk) / 4.0 * 100
    ITI_class = map_matrix_to_report_class(get_matrix_classification("ITI"), "ITI")
    
    SAG_risk = dim_stats.get("SAG", {}).get("risk_score_medio", 1.0)
    SAG_score = (5.0 - SAG_risk) / 4.0 * 100
    SAG_class = map_matrix_to_report_class(get_matrix_classification("SAG"), "SAG")
    
    # Satisfacao is mapped to average of protective dimensions
    satisfacao_score = (REL_score + ORG_score + ITI_score) / 3.0
    satisfacao_risk = 5.0 - (satisfacao_score / 100.0 * 4.0)
    from psicossocial.score import classify_risk_score
    try:
        satisfacao_copsoq = classify_risk_score(satisfacao_risk, metodologia.data["score"])
    except Exception:
        satisfacao_copsoq = "Favoravel"
    satisfacao_class = map_copsoq_to_report_class(satisfacao_copsoq)
    
    total_resp = len(processed_survey.respondents)
    
    # VCO item stats
    vco_stats = {}
    for code in ("Q16_VCO_1", "Q16_VCO_2", "Q16_VCO_3", "Q16_VCO_4", "Q16_VCO_5"):
        yes_count = sum(1 for r in processed_survey.respondents if r.vco.get(code, "").lower().strip() == "sim")
        pct = (yes_count / total_resp * 100.0) if total_resp > 0 else 0.0
        vco_stats[code] = pct
        
    has_vco_1 = vco_stats.get("Q16_VCO_1", 0) > 0
    has_vco_2 = vco_stats.get("Q16_VCO_2", 0) > 0
    has_vco_3 = vco_stats.get("Q16_VCO_3", 0) > 0
    has_vco_4 = vco_stats.get("Q16_VCO_4", 0) > 0
    has_vco_5 = vco_stats.get("Q16_VCO_5", 0) > 0
    
    if has_vco_4 or has_vco_5:
        vco_nivel = 5
        vco_class = "Alto Risco"
        vco_nivel_label = "Nível 5 – Tratativa crítica"
    elif has_vco_3:
        vco_nivel = 4
        vco_class = "Alto Risco"
        vco_nivel_label = "Nível 4 – Tratativa multidisciplinar"
    elif has_vco_2:
        vco_nivel = 3
        vco_class = "Médio Risco"
        vco_nivel_label = "Nível 3 – Avaliação preliminar"
    elif has_vco_1:
        vco_nivel = 2
        vco_class = "Baixo Risco"
        vco_nivel_label = "Nível 2 – Registro e monitoramento"
    else:
        vco_nivel = 1
        vco_class = "Baixo Risco"
        vco_nivel_label = "Nível 1 – Sem ocorrência"

    

    

        
    # 2. General metadata variables
    empresa = projeto.empresa or ""
    cnpj = projeto.cnpj or ""
    
    # Define contract name based on Q2 forms responses (most common)
    q2_values = [
        r.characterization.get("Q2", "").strip()
        for r in processed_survey.respondents
        if r.characterization.get("Q2") and r.characterization.get("Q2").strip()
    ]
    if q2_values:
        from collections import Counter
        nome_contrato = Counter(q2_values).most_common(1)[0][0]
    else:
        nome_contrato = projeto.nome or ""
        
    localidade = projeto.localidade or ""
    
    data_aplicacao = projeto.data_aplicacao.strftime('%d/%m/%Y') if projeto.data_aplicacao else ""
    data_emissao = projeto.data_emissao.strftime('%d/%m/%Y') if projeto.data_emissao else ""
    
    periodo_str = ""
    if projeto.periodo_inicio and projeto.periodo_fim:
        periodo_str = f"{projeto.periodo_inicio.strftime('%d/%m/%Y')} a {projeto.periodo_fim.strftime('%d/%m/%Y')}"
    periodo_mes_ano = ""
    if projeto.periodo_inicio and projeto.periodo_fim:
        periodo_mes_ano = f"{projeto.periodo_inicio.strftime('%m/%Y')} a {projeto.periodo_fim.strftime('%m/%Y')}"
    else:
        periodo_mes_ano = "Não informado"
        
    resp_tecnico = projeto.responsavel_tecnico or ""
    
    total_workforce = sum(r["colaboradores_total"] for r in representatividade_summary if r["recorte"] == "Local")
    
    total_elegiveis = 0
    if total_workforce > 0:
        total_elegiveis = total_workforce
    elif headcounts:
        total_elegiveis = sum(headcounts.values())
    else:
        ghe_summary = [r for r in representatividade_summary if r["recorte"] == "GHE" and r["ghe"] != "Todos"]
        for row in ghe_summary:
            total_elegiveis += row["colaboradores_total"]
        
    if total_elegiveis == 0:
        total_elegiveis = total_resp  # fallback
        
    taxa_adesao = (total_resp / total_elegiveis * 100.0) if total_elegiveis > 0 else 100.0

    # 3. Clean up the instruction and template alert paragraphs first
    for p in list(doc.paragraphs):
        p_text_norm = normalize_text(p.text)
        if "graficos de barras" in p_text_norm and "radar" in p_text_norm:
            delete_paragraph(p)
        elif "orientacao ao preenchimento" in p_text_norm:
            delete_paragraph(p)
            
    # Also delete inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    p_text_norm = normalize_text(p.text)
                    if "orientacao ao preenchimento" in p_text_norm:
                        delete_paragraph(p)

    # 4. Fill Table 1 (Identificação) directly row by row
    table_identificacao = None
    for table in doc.tables:
        if len(table.rows) >= 8 and len(table.columns) == 2:
            if "Empresa" in table.rows[0].cells[0].text or "CNPJ" in table.rows[1].cells[0].text:
                table_identificacao = table
                break
                
    if table_identificacao:
        for row in table_identificacao.rows:
            if len(row.cells) == 2:
                c0 = normalize_text(row.cells[0].text)
                if "empresa" in c0:
                    row.cells[1].text = empresa
                elif "cnpj" in c0:
                    row.cells[1].text = cnpj
                elif "contrato" in c0 or "unidade" in c0:
                    row.cells[1].text = nome_contrato
                elif "municipio" in c0 or "estado" in c0:
                    row.cells[1].text = localidade
                elif "aplicacao" in c0:
                    row.cells[1].text = data_aplicacao
                elif "apuracao" in c0:
                    row.cells[1].text = periodo_str if periodo_str else "Não informado"
                elif "responsavel" in c0:
                    row.cells[1].text = resp_tecnico
                elif "emissao" in c0:
                    row.cells[1].text = data_emissao

    # 5. Fill Table 3 (Metodologia)
    table_metodologia = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 2:
            first_cell_norm = normalize_text(table.rows[0].cells[0].text)
            if "aspecto metodologico" in first_cell_norm:
                table_metodologia = table
                break
                
    if table_metodologia:
        # Mantido estático conforme o novo modelo padrão
        pass

    # 6. Fill Table 4 (Caracterização da Amostra)
    table_amostra = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 2:
            first_cell_norm = normalize_text(table.rows[1].cells[0].text)
            if "trabalhadores elegiveis" in first_cell_norm:
                table_amostra = table
                break
                
    if table_amostra:
        for row in table_amostra.rows[1:]:
            if len(row.cells) == 2:
                header = normalize_text(row.cells[0].text)
                if "trabalhadores elegiveis" in header or "elegiveis GHE" in header:
                    row.cells[1].text = f"{total_elegiveis} colaboradores"
                elif "participantes" in header:
                    row.cells[1].text = f"{total_resp} colaboradores"
                elif "taxa de adesao" in header:
                    row.cells[1].text = f"{taxa_adesao:.1f}% — adequada para representatividade" if taxa_adesao >= 60.0 else f"{taxa_adesao:.1f}%"
                elif "criterio de inclusao" in header:
                    row.cells[1].text = "Todos os colaboradores ativos no cadastro SRA."
                elif "criterio de exclusao" in header:
                    row.cells[1].text = "Colaboradores desligados no período ou ausentes."
                elif "representatividade" in header:
                    if taxa_adesao >= 60.0:
                        row.cells[1].text = (
                            f"Trabalhadores elegíveis GHE: {total_elegiveis} colaboradores | "
                            f"Participantes (respostas válidas): {total_resp} colaboradores | "
                            f"Taxa de adesão: {taxa_adesao:.1f}% — adequada para representatividade"
                        )
                    else:
                        row.cells[1].text = (
                            f"Trabalhadores elegíveis GHE: {total_elegiveis} colaboradores | "
                            f"Participantes (respostas válidas): {total_resp} colaboradores | "
                            f"Taxa de adesão: {taxa_adesao:.1f}% — representatividade amostral parcial."
                        )

    # 7. Fill Table 5 (Resultados Consolidados)
    table_consolidados = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) >= 3:
            first_row_text = [normalize_text(cell.text) for cell in table.rows[0].cells]
            if any("dimensao" in txt for txt in first_row_text) and any("classificacao" in txt for txt in first_row_text):
                table_consolidados = table
                break
                
    if table_consolidados:
        rows_to_delete = []
        for row in table_consolidados.rows:
            if len(row.cells) >= 3:
                cell_0_norm = normalize_text(row.cells[0].text)
                if "satisfacao" in cell_0_norm:
                    rows_to_delete.append(row)
                    continue
                # Determine if this is the score column (censor VCO only, show other scores)
                def _set_score_row(name_cell, score_cell, class_cell, name_text, score, classification):
                    name_cell.text = name_text
                    if score is None:
                        score_cell.text = "**"
                    else:
                        score_cell.text = f"{score:.1f}%"
                    style_class_cell(class_cell, classification)
                if "demandas" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Demandas de Trabalho (DEM)", DEM_score, DEM_class)
                elif "estresse" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Estresse e Esgotamento (EST)", EST_score, EST_class)
                elif "conflito" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Conflito Trabalho-Vida (CTV)", CTV_score, CTV_class)
                elif "relacoes" in cell_0_norm or "lideranca" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Relações Liderança (REL)", REL_score, REL_class)
                elif "organizacao" in cell_0_norm or "autonomia" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Organização e Autonomia (ORG)", ORG_score, ORG_class)
                elif "interface" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Interface Trabalho-Indivíduo (ITI)", ITI_score, ITI_class)
                elif "saude" in cell_0_norm:
                    _set_score_row(row.cells[0], row.cells[1], row.cells[2], "Saúde Geral e Bem-estar (SAG)", SAG_score, SAG_class)
                elif "violencia" in cell_0_norm or "ofensivos" in cell_0_norm:
                    row.cells[0].text = "Violência e Comportamentos ofensivos (> identificado)"
                    row.cells[1].text = "**"
                    style_class_cell(row.cells[2], vco_class)
        for r in rows_to_delete:
            table_consolidados._tbl.remove(r._tr)

    # 8. Fill Table 7 (VCO)
    table_vco = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 2:
            if "Tipo de Comportamento Ofensivo" in table.rows[0].cells[0].text:
                table_vco = table
                break
                
    if table_vco:
        # Pre-calculate combination for Q16_VCO_2 and Q16_VCO_5
        yes_2_or_5_count = sum(1 for r in processed_survey.respondents if r.vco.get("Q16_VCO_2", "").lower().strip() == "sim" or r.vco.get("Q16_VCO_5", "").lower().strip() == "sim")
        vco_2_or_5_pct = (yes_2_or_5_count / total_resp * 100.0) if total_resp > 0 else 0.0

        for row in table_vco.rows:
            if len(row.cells) == 2:
                header = normalize_text(row.cells[0].text)
                if "assedio verbal" in header or "humilhacao" in header:
                    val_pct = vco_stats["Q16_VCO_2"]
                    row.cells[1].text = f"{val_pct:.1f}% dos respondentes" if val_pct > 0 else "Não registrado (0%)"
                elif "desrespeitoso" in header or "superiores" in header:
                    val_pct = vco_stats["Q16_VCO_1"]
                    row.cells[1].text = f"{val_pct:.1f}% dos respondentes" if val_pct > 0 else "Não registrado (0%)"
                elif "discriminacao" in header:
                    row.cells[1].text = "Não registrado (0%)"
                elif "assedio moral" in header or "moral" in header:
                    val_pct = vco_stats["Q16_VCO_3"]
                    row.cells[1].text = f"{val_pct:.1f}% dos respondentes" if val_pct > 0 else "Não registrado (0%)"
                elif "ameacas" in header or "fisica" in header:
                    val_pct = vco_2_or_5_pct
                    row.cells[1].text = f"{val_pct:.1f}% dos respondentes" if val_pct > 0 else "Não registrado (0%)"
                elif "sexual" in header:
                    val_pct = vco_stats["Q16_VCO_4"]
                    row.cells[1].text = f"{val_pct:.1f}% dos respondentes" if val_pct > 0 else "Não registrado (0%)"

    # 9. Fill Table 8 (Indicadores Organizacionais)
    table_indicadores = None
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) == 3:
            first_cell_norm = normalize_text(table.rows[1].cells[0].text)
            if "absenteismo" in first_cell_norm:
                table_indicadores = table
                break
                
    if table_indicadores:
        for row in table_indicadores.rows[1:]:
            if len(row.cells) == 3:
                row.cells[1].text = periodo_mes_ano
                header = normalize_text(row.cells[0].text)
                if "absenteismo" in header:
                    row.cells[2].text = "0.0%"
                elif "turnover" in header:
                    row.cells[2].text = "0.0%"
                elif "afastamentos" in header:
                    row.cells[2].text = "Nenhum afastamento registrado"
                elif "canal de etica" in header or "registros canal" in header:
                    row.cells[2].text = "Nenhum registro formal"
                elif "consultas" in header or "atendimentos" in header or "pcmso" in header:
                    row.cells[2].text = "Nenhum atendimento registrado"

    # 10. Dynamic Table 4.1 (Distribuição por GHE / Função / Setor) + Setor distribution
    table_representatividade = None
    template_tr = None
    for table in doc.tables:
        if len(table.rows) > 0:
            first_row_text = [cell.text for cell in table.rows[0].cells]
            if any("Elegíveis" in txt or "Elegveis" in txt or "GHE" in txt for txt in first_row_text) and any("Adesão" in txt or "Adeso" in txt for txt in first_row_text) and not any("Dimensão" in txt or "Nivel de Risco" in txt for txt in first_row_text):
                table_representatividade = table
                break

    if table_representatividade and len(table_representatividade.rows) > 1:
        tbl = table_representatividade._tbl
        # Save template row XML from the first mock data row
        template_tr = copy.deepcopy(table_representatividade.rows[1]._tr)

        # Delete all mock rows except the header
        while len(table_representatividade.rows) > 1:
            tbl.remove(table_representatividade.rows[1]._tr)

        # ── Setor/Área distribution (Q1 x Q6) ──────────────
        from collections import Counter as _Counter
        setor_counts = _Counter()
        for resp in processed_survey.respondents:
            q1_val = resp.characterization.get("Q1", "").strip()
            q6_raw = resp.characterization.get("Q6", "").strip()
            q6_val = normalize_activity_q6(q6_raw)
            if q1_val and q6_val:
                key = f"{q1_val} | {q6_val}"
                setor_counts[key] += 1
            else:
                key = q1_val or q6_val or "Não informado"
                setor_counts[key] += 1

        # Add a column to the header row XML if it has only 4 columns
        header_tr = table_representatividade.rows[0]._tr
        tcs_header = header_tr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
        if tcs_header and len(tcs_header) == 4:
            new_tc = copy.deepcopy(tcs_header[-1])
            header_tr.append(new_tc)

        # Add a column to the template row XML if it has only 4 columns
        tcs_template = template_tr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
        if tcs_template and len(tcs_template) == 4:
            new_tc = copy.deepcopy(tcs_template[-1])
            template_tr.append(new_tc)

        # Override header titles to ensure correct layout
        header_cells = table_representatividade.rows[0].cells
        if len(header_cells) == 6:
            header_cells[0].text = "GHE / Setor Q1"
            header_cells[1].text = "Tipo de atividade Q6"
            header_cells[2].text = "Total"
            header_cells[3].text = "Elegíveis"
            header_cells[4].text = "Adesão"
            header_cells[5].text = "Interpretação"
            align_indices = (2, 3, 4, 5)
        elif len(header_cells) >= 5:
            header_cells[0].text = "GHE / Setor"
            header_cells[1].text = "Total"
            header_cells[2].text = "Elegíveis"
            header_cells[3].text = "Adesão"
            header_cells[4].text = "Interpretação"
            align_indices = (1, 2, 3, 4)
        else:
            align_indices = list(range(1, len(header_cells)))

        for col_idx in align_indices:
            if col_idx < len(header_cells):
                for p in header_cells[col_idx].paragraphs:
                    p.alignment = 1

        # One row per combination, sorted by count descending
        norm_headcounts = {normalize_headcount_key(k): v for k, v in (headcounts or {}).items()}

        # Include GHE/area combinations with an informed headcount but zero survey
        # respondents (e.g. the "Não identificado" bucket for the contract's total
        # headcount vs. the areas covered by responses), so non-respondents are
        # visible in the report instead of only affecting the aggregate adherence rate.
        existing_norm_keys = {normalize_headcount_key(k) for k in setor_counts}
        for raw_key in (headcounts or {}):
            norm_key = normalize_headcount_key(raw_key)
            if norm_key not in existing_norm_keys:
                setor_counts[raw_key] = 0
                existing_norm_keys.add(norm_key)

        for key, count in sorted(setor_counts.items(), key=lambda x: -x[1]):
            if " | " in key:
                q1_val, q6_val = key.split(" | ", 1)
            else:
                q1_val, q6_val = key, ""
                
            norm_key = normalize_headcount_key(key)
            total_val = norm_headcounts.get(norm_key, 0)
            if total_val > 0:
                adesao_pct = (count / total_val) * 100.0
                if adesao_pct > 100.0:
                    adesao_pct = 100.0
                adesao_str = f"{adesao_pct:.1f}%"
                total_str = str(total_val)
                
                if adesao_pct < 40.0:
                    interpretation = "Baixa confiabilidade - Não divulgar resultados; reforçar estratégia de comunicação."
                elif adesao_pct < 60.0:
                    interpretation = "Confiabilidade moderada - Análise restrita; uso apenas interno e técnico."
                elif adesao_pct < 80.0:
                    interpretation = "Boa representatividade - Divulgação permitida com registro de limitação."
                else:
                    interpretation = "Alta representatividade - Divulgação plena com alta confiança metodológica."
            else:
                adesao_str = "**"
                total_str = "**"
                interpretation = ""
                
            row_cells = add_row_from_xml(table_representatividade, template_tr).cells
            if len(row_cells) == 6:
                row_cells[0].text = q1_val
                row_cells[1].text = q6_val
                row_cells[2].text = total_str
                row_cells[3].text = str(count)
                row_cells[4].text = adesao_str
                row_cells[5].text = interpretation
                for idx in (2, 3, 4, 5):
                    for p in row_cells[idx].paragraphs:
                        p.alignment = 1
            elif len(row_cells) >= 5:
                row_cells[0].text = q1_val
                row_cells[1].text = total_str
                row_cells[2].text = str(count)
                row_cells[3].text = adesao_str
                row_cells[4].text = interpretation
                for idx in (1, 2, 3, 4):
                    for p in row_cells[idx].paragraphs:
                        p.alignment = 1
            else:
                row_cells[0].text = q1_val
                if len(row_cells) > 1:
                    row_cells[1].text = total_str
                if len(row_cells) > 2:
                    row_cells[2].text = str(count)
                if len(row_cells) > 3:
                    row_cells[3].text = adesao_str
                for idx in range(1, len(row_cells)):
                    for p in row_cells[idx].paragraphs:
                        p.alignment = 1

    # 11. Dynamic Table 10 (Classificação Final) using template XML
    from psicossocial.processing import dimension_exposure_summary
    from psicossocial.workforce import canonical_ghe
    
    ghe_respondents = {}
    for r in processed_survey.respondents:
        ghe = canonical_ghe(r.characterization.get("Q1", ""))
        if not ghe or ghe.lower() in ("não informado", "nao informado"):
            ghe = "Não informado"
        ghe_respondents.setdefault(ghe, []).append(r)
            
    critical_exposures = []
    dim_mappings = {
        "DEM": ("Demandas de Trabalho", "Pressão e Sobrecarga"),
        "EST": ("Estresse e Esgotamento", "Desgaste Emocional"),
        "CTV": ("Conflito Trabalho x Vida", "Conflito Trabalho-Família"),
        "REL": ("Relações e Liderança", "Suporte Liderança/Colegas"),
        "ORG": ("Organização e Autonomia", "Autonomia e Fluxo de Tarefas"),
        "ITI": ("Interface Trabalho-Indivíduo", "Reconhecimento/Expectativas"),
        "SAG": ("Saúde Geral", "Saúde Física e Mental"),
    }
    
    for ghe, resp_list in sorted(ghe_respondents.items()):
        for dim_code in ("DEM", "EST", "CTV", "REL", "ORG", "ITI", "SAG"):
            try:
                exposure = dimension_exposure_summary(tuple(resp_list), dim_code)
                if exposure.matrix_risk_classification in ("Alto", "Critico"):
                    dim_name, factor_name = dim_mappings[dim_code]
                    critical_exposures.append({
                        "ghe": ghe,
                        "dim": dim_code,
                        "fator": factor_name,
                        "prob": exposure.probability_level,
                        "sev": exposure.severity_level,
                        "nivel": f"{exposure.matrix_risk_level} - {exposure.matrix_risk_classification.upper()}"
                    })
            except Exception:
                continue
                
    table_classificacao = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 6:
            first_row_text = [cell.text for cell in table.rows[0].cells]
            if any("ghe" in normalize_text(txt) for txt in first_row_text) and any("dimens" in normalize_text(txt) for txt in first_row_text) and any("nivel de risco" in normalize_text(txt) for txt in first_row_text):
                table_classificacao = table
                break
                
    if table_classificacao and len(table_classificacao.rows) > 1:
        # Save template row XML
        template_tr = copy.deepcopy(table_classificacao.rows[1]._tr)
        
        # Delete all mock rows except header
        while len(table_classificacao.rows) > 1:
            table_classificacao._tbl.remove(table_classificacao.rows[1]._tr)
        
        if critical_exposures:
            for exp in critical_exposures:
                row_cells = add_row_from_xml(table_classificacao, template_tr).cells
                row_cells[0].text = exp["ghe"]
                row_cells[1].text = exp["dim"]
                row_cells[2].text = exp["fator"]
                row_cells[3].text = str(exp["prob"])
                row_cells[4].text = str(exp["sev"])
                style_class_cell(row_cells[5], exp["nivel"])
        else:
            row = add_row_from_xml(table_classificacao, template_tr)
            row_cells = row.cells
            row_cells[0].text = "Nenhum risco psicossocial Alto ou Crítico foi identificado nos GHEs avaliados."
            row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4]).merge(row_cells[5])

    # 12. Dynamic Table 11 (Plano de Ação) using template XML
    actions_to_add = []
    action_templates = {
        "DEM": {
            "fator": "DEM - Sobrecarga e ritmo de trabalho",
            "class": DEM_class,
            "acao": "Revisão do dimensionamento de equipes, distribuição de tarefas e definição clara de metas para reduzir a pressão temporal.",
            "resp": "Gestor da área / SSMA",
            "prazo": "Até 90 dias",
            "monit": "Reaplicação COPSOQ em 12 meses; indicadores de absenteísmo e turnover."
        },
        "EST": {
            "fator": "EST - Desgaste e exaustão emocional",
            "class": EST_class,
            "acao": "Implementação de programas de promotion à saúde mental, treinamentos de resiliência e suporte psicossocial aos colaboradores.",
            "resp": "RH / Gestão de Contratos",
            "prazo": "Até 90 dias",
            "monit": "NPS interno; indicadores de afastamentos médicos por CID F."
        },
        "CTV": {
            "fator": "CTV - Conflito entre trabalho e vida pessoal",
            "class": CTV_class,
            "acao": "Ações de flexibilização de jornada quando viável, respeito aos períodos de descanso e campanhas de integração família-trabalho.",
            "resp": "RH / Gestores de Área",
            "prazo": "Até 90 dias",
            "monit": "Pesquisa de satisfação interna e acompanhamento de escalas."
        },
        "REL": {
            "fator": "REL - Relações socioprofissionais e liderança",
            "class": REL_class,
            "acao": "Treinamento de liderança positiva, comunicação assertiva e canais formais de feedback entre liderança e equipes.",
            "resp": "RH / Liderança de Contrato",
            "prazo": "Até 60 dias",
            "monit": "Pesquisa de clima anual e monitoramento de conflitos internos."
        },
        "ORG": {
            "fator": "ORG - Organização das tarefas e autonomia",
            "class": ORG_class,
            "acao": "Revisão e melhoria do fluxo de informações operacionais, estímulo à participação em sugestões de melhoria.",
            "resp": "Gestor do Contrato / CIPA",
            "prazo": "Até 90 dias",
            "monit": "Sugestões operacionais acolhidas e implantadas."
        },
        "ITI": {
            "fator": "ITI - Reconhecimento e alinhamento de expectativas",
            "class": ITI_class,
            "acao": "Alinhamento de expectativas no onboarding, programas de reconhecimento por desempenho e aproveitamento de habilidades internas.",
            "resp": "RH / Liderança de Contrato",
            "prazo": "Até 90 dias",
            "monit": "Pesquisa de engajamento e taxa de retenção."
        },
        "SAG": {
            "fator": "SAG - Percepção de saúde física e mental",
            "class": SAG_class,
            "acao": "Campanhas preventivas de saúde física (ergonomia, ginástica laboral) e mental, em parceria com o PCMSO.",
            "resp": "Médico do Trabalho / SESMT",
            "prazo": "Até 60 dias",
            "monit": "Indicadores do PCMSO e acompanhamento de queixas."
        },
        "VCO": {
            "fator": "VCO - Comportamentos ofensivos (Assédio / Discriminação)",
            "class": vco_class,
            "acao": "Divulgação ostensiva do Canal de Ética, treinamento obrigatório sobre prevenção ao assédio e fluxo de tratativas rápidas.",
            "resp": "Comitê de Ética / RH / CIPA",
            "prazo": "Imediato / Até 30 dias",
            "monit": "Treinamentos realizados (adesão %) e denúncias resolvidas."
        }
    }
    
    for dim_code, data in action_templates.items():
        if dim_code == "VCO":
            continue
        cls_val = data["class"]
        if "Alto" in cls_val or "Médio" in cls_val or "Critico" in cls_val or "Grave" in cls_val:
            actions_to_add.append(data)
            
    if not actions_to_add:
        actions_to_add.append({
            "fator": "Monitoramento Preventivo Geral",
            "class": "Baixo Risco",
            "acao": "Manter os fatores protetivos ativos através de reuniões regulares de feedback e acompanhamento contínuo dos canais de escuta.",
            "resp": "Liderança / RH",
            "prazo": "Contínuo",
            "monit": "Reaplicação anual da avaliação psicossocial."
        })
        
    table_plano = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 6:
            first_row_text = [cell.text for cell in table.rows[0].cells]
            if any("fator psicossocial" in normalize_text(txt) for txt in first_row_text) and any("monitoramento" in normalize_text(txt) for txt in first_row_text):
                table_plano = table
                break
                
    if table_plano and len(table_plano.rows) > 1:
        # Save template row XML
        template_tr = copy.deepcopy(table_plano.rows[1]._tr)
        
        # Delete all mock rows except header
        while len(table_plano.rows) > 1:
            table_plano._tbl.remove(table_plano.rows[1]._tr)
        
        for act in actions_to_add:
            row_cells = add_row_from_xml(table_plano, template_tr).cells
            row_cells[0].text = act["fator"]
            style_class_cell(row_cells[1], act["class"])
            row_cells[2].text = act["acao"]
            row_cells[3].text = act["resp"]
            row_cells[4].text = act["prazo"]
            row_cells[5].text = act["monit"]

    # 13. Fill Table 12 (Integração com AEP, PGR, PCMSO)
    table_integracao = None
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) == 3:
            if "PGR" in table.rows[1].cells[0].text and "PCMSO" in table.rows[2].cells[0].text:
                table_integracao = table
                break
                
    if table_integracao:
        for row in table_integracao.rows[1:]:
            if len(row.cells) == 3:
                header = normalize_text(row.cells[0].text)
                if "pgr" in header or "inventario" in header:
                    row.cells[2].text = "Revisar PGR com a inclusão dos fatores de risco psicossocial identificados e suas respectivas medidas de controle."
                elif "pcmso" in header:
                    row.cells[2].text = "Considerar a avaliação psicossocial, histórico de queixas e absenteísmo para o planejamento de ações preventivas de saúde no PCMSO."
                elif "monitoramento" in header:
                    row.cells[2].text = "Reaplicação da avaliação em 12 meses; monitoramento do plano de ação em 6 meses."

    # 14. Fill Signatures Table directly
    table_assinaturas = None
    for table in doc.tables:
        if len(table.rows) == 2 and len(table.columns) == 2:
            if "Elaboração" in table.rows[0].cells[0].text or "Elaboracao" in table.rows[0].cells[0].text:
                table_assinaturas = table
                break
                
    if table_assinaturas and len(table_assinaturas.rows) == 2:
        c0 = table_assinaturas.rows[1].cells[0]
        c0.text = ""
        p0 = c0.paragraphs[0]
        p0.add_run(f"{resp_tecnico}\nResponsável Técnico SSMA\nData: {data_emissao if data_emissao else '____/____/____'}\nAssinatura: _________________________")
        
        c1 = table_assinaturas.rows[1].cells[1]
        c1.text = ""
        p1 = c1.paragraphs[0]
        p1.add_run(f"Gerente SSMA / Médico Coordenador PCMSO\nData: {data_emissao if data_emissao else '____/____/____'}\nAssinatura: _________________________")

    # 15. Setup descriptions and build lists of risks vs protection factors
    DIM_DESCRIPTIONS = {
        "DEM": "relacionada ao ritmo laboral e esforço cognitivo.",
        "EST": "relacionado ao estresse, desgaste mental e cansaço decorrente da atividade.",
        "CTV": "reflexo do equilíbrio entre vida pessoal e obrigações profissionais.",
        "REL": "cooperação mútua entre equipes e liderança direta.",
        "ORG": "autonomia de tarefas, clareza de papel e fluxo organizacional.",
        "ITI": "alinhamento de expectativas, reconhecimento profissional e perspectivas de futuro.",
        "SAG": "percepção de saúde física, mental e bem-estar geral.",
        "satisfacao": "contentamento geral do trabalhador com a sua função e condições de trabalho.",
        "VCO": "registros consolidados das respostas declaratórias do Bloco VCO."
    }

    DIM_SHORT_NAMES = {
        "DEM": "Demandas de Trabalho",
        "EST": "Estresse e Esgotamento",
        "CTV": "Conflito Trabalho x Vida",
        "REL": "Relações e Liderança",
        "ORG": "Organização e Autonomia",
        "ITI": "Interface Trabalho-Indivíduo",
        "SAG": "Saúde Geral",
        "satisfacao": "Satisfação com o Trabalho",
        "VCO": "Comportamentos Ofensivos",
    }

    dimensions_data = [
        ("DEM", "Demandas de Trabalho", DEM_score, DEM_class),
        ("EST", "Estresse e Esgotamento", EST_score, EST_class),
        ("CTV", "Conflito Trabalho x Vida", CTV_score, CTV_class),
        ("REL", "Relações e Liderança", REL_score, REL_class),
        ("ORG", "Organização e Autonomia", ORG_score, ORG_class),
        ("ITI", "Interface Trabalho-Indivíduo", ITI_score, ITI_class),
        ("SAG", "Saúde Geral", SAG_score, SAG_class),
        # "satisfacao" removido conforme padrão do relatório
    ]

    def get_prevalence(code):
        if code == "VCO":
            exposed = sum(1 for r in processed_survey.respondents if any(v == "sim" for v in r.vco.values()))
            return (exposed / total_resp * 100.0) if total_resp > 0 else 0.0
        scores = [
            r.dimensions[code].risk_score
            for r in processed_survey.respondents
            if code in r.dimensions
        ]
        exposed_count = sum(1 for s in scores if s >= 4.0)
        return (exposed_count / len(scores) * 100.0) if scores else 0.0

    # Determine which dimensions go to 6.2 (Dimensões com Maior Criticidade)
    # Target: ONLY Alto Risco (no fallbacks to Médio or Baixo Risco)
    crit_candidates = [d for d in dimensions_data if d[3] == "Alto Risco"]

    # VCO Candidate for 6.2 (desativado conforme novo padrão de isolar VCO da seção 6.2)
    vco_in_6_2 = False

    risk_items = []
    for code, name, score, classification in crit_candidates:
        item_desc, ghe_desc = analyze_dimension_criticality(processed_survey, code, metodologia)
        item_str = f"{name} (prevalência {score:.1f}% – {classification}): relato frequente de {item_desc}, especialmente relatado por trabalhadores do {ghe_desc}."
        risk_items.append(item_str)

    if vco_in_6_2:
        vco_item_desc, vco_ghe_desc = analyze_vco_criticality(processed_survey)
        vco_item_str = f"Violência e Comportamentos Ofensivos (prevalência ** – {vco_class}): registro de situações de {vco_item_desc}, especialmente relatadas por trabalhadores do {vco_ghe_desc}."
        risk_items.append(vco_item_str)

    # Determine which dimensions go to 6.3 (Fatores Protetivos Identificados)
    # Show dimensions with a Baixo Risco or Fator Protetor classification.
    prot_items = []
    for code, name, score, classification in dimensions_data:
        if classification in ("Baixo Risco", "Fator Protetor"):
            desc = DIM_DESCRIPTIONS[code]
            prot_items.append(f"{name} (escore {score:.1f} – {classification}): {desc}")

    # 17. Narratives for 6.1, 6.4, and Section 12 (Process 6.1 first to avoid placeholder collisions)
    high_risk_dims = [DIM_SHORT_NAMES[code] for code, name, score, classification in dimensions_data if classification == "Alto Risco"]
    if vco_class == "Alto Risco":
        high_risk_dims.append(DIM_SHORT_NAMES["VCO"])
        
    med_risk_dims = [DIM_SHORT_NAMES[code] for code, name, score, classification in dimensions_data if classification == "Médio Risco"]
    if vco_class == "Médio Risco":
        med_risk_dims.append(DIM_SHORT_NAMES["VCO"])
        
    if high_risk_dims:
        summary_riscos = f"As dimensões {', '.join(high_risk_dims)} apresentaram score de maior criticidade (Alto Risco), indicando áreas de atenção prioritária no ambiente de trabalho."
    elif med_risk_dims:
        summary_riscos = f"As dimensões {', '.join(med_risk_dims)} apresentaram criticidade moderada (Médio Risco), necessitando de ações preventivas específicas."
    else:
        summary_riscos = "Nenhuma dimensão apresentou score de alto ou médio risco crítico. O cenário geral demonstra fatores de risco controlados."

    sec_6_1_paras = [
        summary_riscos,
        f"Demandas de Trabalho (escore {DEM_score:.1f} – {DEM_class}): relacionada ao ritmo laboral e espaço cognitivo.",
        f"Estresse e Esgotamento (escore {EST_score:.1f} – {EST_class}): relacionado ao estresse, desgaste mental e cansaço decorrente da atividade.",
        f"Conflito Trabalho x Vida (escore {CTV_score:.1f} – {CTV_class}): reflexo do equilíbrio entre vida pessoal e obrigações profissionais.",
        f"Relações e Liderança (escore {REL_score:.1f} – {REL_class}): cooperação mútua entre equipes e liderança direta.",
        f"Organização e Autonomia (escore {ORG_score:.1f} – {ORG_class}): autonomia de tarefas, clareza de papel e fluxo organizacional.",
        f"Interface Trabalho-Indivíduo (escore {ITI_score:.1f} – {ITI_class}): alinhamento de expectativas, reconhecimento profissional e perspectivas de futuro.",
        f"Saúde Geral (escore {SAG_score:.1f} – {SAG_class}): percepção de saúde física, mental e bem-estar geral."
    ]
    replace_heading_with_multiple_paragraphs(doc, "6.1", sec_6_1_paras)

    # 16. Dyn write Section 6.2 and 6.3 bullet lists
    populate_section_bullets(doc, "6.2", risk_items)
    populate_section_bullets(doc, "6.3", prot_items)
    
    # Append the user's custom protective factors text input at the end of Section 6.3
    if hasattr(projeto, "detalhamento_fatores_protetivos") and projeto.detalhamento_fatores_protetivos:
        append_free_text_to_section_6_3(doc, projeto.detalhamento_fatores_protetivos)
    
    tendencias_text = "A combinação de eventuais scores desfavoráveis (Médio ou Alto Risco) indica áreas que demandam monitoramento para prevenção de transtornos ansiosos ou esgotamento laboral. O plano de ação preventivo mitigará potenciais impactos nos indicadores de absenteísmo da contratada."
    replace_heading_example(doc, "6.4", tendencias_text)

    # Technical Conclusion (Section 12) paragraphs
    num_high = len(high_risk_dims)
    num_med = len(med_risk_dims)
    
    def pt_number_words(n):
        words = ["zero", "uma", "duas", "três", "quatro", "cinco", "seis", "sete", "oito", "nove", "dez"]
        if 0 <= n < len(words):
            return words[n]
        return str(n)

    if num_high > 0 or num_med > 0:
        high_part = ""
        if num_high > 0:
            high_list_str = ", ".join(high_risk_dims[:-1]) + (" e " + high_risk_dims[-1] if len(high_risk_dims) > 1 else high_risk_dims[0])
            high_part = f"{pt_number_words(num_high)} dimensões em Alto Risco – {high_list_str} –"
            
        med_part = ""
        if num_med > 0:
            med_list_str = ", ".join(med_risk_dims[:-1]) + (" e " + med_risk_dims[-1] if len(med_risk_dims) > 1 else med_risk_dims[0])
            med_part = f"{pt_number_words(num_med)} dimensões em Médio Risco – {med_list_str} –"
            
        parts = [p for p in (high_part, med_part) if p]
        new_conclusion_1 = f"A avaliação psicossocial realizada na unidade {nome_contrato} identificou " + ", além de ".join(parts).replace(" –,", " –") + "."
    else:
        new_conclusion_1 = f"A avaliação psicossocial realizada na unidade {nome_contrato} não identificou nenhuma dimensão em Alto ou Médio Risco, demonstrando um perfil favorável de saúde mental e clima psicossocial."
        
    prot_names_short = [DIM_SHORT_NAMES[code] for code, name, score, classification in dimensions_data if classification == "Baixo Risco"]
    if vco_class == "Baixo Risco":
        prot_names_short.append(DIM_SHORT_NAMES["VCO"])
        
    if prot_names_short:
        prot_list_str = ", ".join(prot_names_short[:-1]) + (" e " + prot_names_short[-1] if len(prot_names_short) > 1 else prot_names_short[0])
        new_conclusion_2 = f"Os fatores protetivos identificados, especialmente {prot_list_str}, representam uma base positiva para a sustentabilidade do clima psicossocial e a implementação das ações propostas."
    else:
        new_conclusion_2 = "Não foram identificados fatores protetivos expressivos, indicando a necessidade de fortalecimento geral das relações de liderança e suporte social."
        
    next_year_date = (datetime.date.today() + datetime.timedelta(days=365)).strftime('%m/%Y')
    new_conclusion_3 = f"Recomenda-se o acompanhamento do plano de ação preventivo com foco nas dimensões críticas identificadas. A reaplicação completa desta avaliação psicossocial está prevista para daqui a 12 meses ({next_year_date})."

    # Write Section 12 paragraphs directly
    heading_12_p = None
    for p in doc.paragraphs:
        if "conclusao tecnica" in normalize_text(p.text):
            heading_12_p = p
            break
            
    if heading_12_p:
        sec_12_paras = []
        current = heading_12_p._p.getnext()
        from docx.text.paragraph import Paragraph
        while current is not None:
            p = Paragraph(current, heading_12_p._parent)
            p_text_norm = normalize_text(p.text)
            if re.match(r"^\d+(\.\d+)*\s+", p.text.strip()) or p.style.name.startswith("Heading"):
                break
            sec_12_paras.append(p)
            current = current.getnext()
            
        conclusions = [new_conclusion_1, new_conclusion_2, new_conclusion_3]
        for i, text in enumerate(conclusions):
            if i < len(sec_12_paras):
                sec_12_paras[i].text = text
            else:
                prev_p = sec_12_paras[i-1] if i > 0 else heading_12_p
                new_p_xml = OxmlElement('w:p')
                prev_p._element.addnext(new_p_xml)
                new_p = Paragraph(new_p_xml, heading_12_p._parent)
                new_p.text = text
                if sec_12_paras:
                    copy_paragraph_format(sec_12_paras[0], new_p)
                sec_12_paras.append(new_p)
                
        if len(sec_12_paras) > len(conclusions):
            for p in sec_12_paras[len(conclusions):]:
                delete_paragraph(p)

    # 18. Fix heading numbering typo in template
    for p in doc.paragraphs:
        p_text_stripped = p.text.strip()
        if p_text_stripped.startswith("13. GOVERNANÇA, CONFIDENCIALIDADE") or p_text_stripped.startswith("13. GOVERNANCA, CONFIDENCIALIDADE"):
            p.text = p.text.replace("13. GOVERNANÇA", "14. GOVERNANÇA").replace("13. GOVERNANCA", "14. GOVERNANÇA")
        elif p_text_stripped.startswith("14. REFERÊNCIAS TÉCNICAS") or p_text_stripped.startswith("14. REFERENCIAS TECNICAS"):
            p.text = p.text.replace("14. REFERÊNCIAS", "15. REFERÊNCIAS").replace("14. REFERENCIAS", "15. REFERÊNCIAS")

    # 19. Clean up VCO POP flow alert below Table 11
    for p in list(doc.paragraphs):
        p_text_norm = normalize_text(p.text)
        if "fluxo do pop" in p_text_norm and "tratativa vco" in p_text_norm:
            if vco_class in ("Alto Risco", "Médio Risco"):
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                delete_paragraph(p)

    # 20. Replace the indicators description typo
    for p in doc.paragraphs:
        p_text_norm = normalize_text(p.text)
        if "permanecerao em monitoramento" in p_text_norm or "absenteismo de 4,2%" in p_text_norm:
            p.text = "Os indicadores organizacionais no período de apuração permanecem sob controle e em monitoramento. Recomenda-se o acompanhamento contínuo dos indicadores de absenteísmo e afastamentos para prevenção e mitigação de impactos relacionados a riscos psicossociais."

    # Direct fallback scan for section 6.4 and 6.1 example paragraphs
    for p in doc.paragraphs:
        p_text_norm = normalize_text(p.text)
        if "combinacao de alta demanda" in p_text_norm or ("tendencias" in p_text_norm and "ex.:" in p_text_norm):
            p.text = tendencias_text
        if "principais fatores de risco" in p_text_norm and ("descreva" in p_text_norm or "ex.:" in p_text_norm):
            p.text = summary_riscos

    # 21. Fallback search & replace for minor standalone metadata items
    replace_regex_everywhere(doc, r"\[Nome completo da empresa\s*/\s*contratada\]", empresa)
    replace_regex_everywhere(doc, r"\[00\.000\.000/0000-00\]", cnpj)
    replace_regex_everywhere(doc, r"\[Nº do contrato\s*[–-]\s*Nome do estabelecimento ou unidade avaliada\]", nome_contrato)
    replace_regex_everywhere(doc, r"\[Nome da Unidade\s*/\s*Contrato\]", nome_contrato)
    replace_regex_everywhere(doc, r"\[Cidade\s*[–-]\s*UF\]", localidade)
    replace_regex_everywhere(doc, r"\[Nome do contrato ou unidade\]", nome_contrato)
    
    # Substituir os novos placeholders qualitativos de VCO e [Nome do contrato ou unidade] adicionais
    exposed_vco_count = sum(1 for r in processed_survey.respondents if any(v == "sim" for v in r.vco.values()))
    vco_existencia = "Sim" if exposed_vco_count > 0 else "Não"
    if exposed_vco_count == 0:
        vco_qtd = "0"
    elif exposed_vco_count < 5:
        vco_qtd = "1 a 4"
    else:
        vco_qtd = str(exposed_vco_count)

    for p in doc.paragraphs:
        p_text_norm = normalize_text(p.text)
        if "bloco vco" in p_text_norm:
            p.text = p.text.replace("[número]", str(total_resp)).replace("[numero]", str(total_resp))
        if "existencia de resposta positiva" in p_text_norm:
            p.text = p.text.replace("[Sim/Não]", vco_existencia).replace("[Sim/Nao]", vco_existencia)
        if "quantidade consolidada de respostas" in p_text_norm:
            p.text = p.text.replace("[número ou faixa, conforme anonimato]", vco_qtd).replace("[numero ou faixa, conforme anonimato]", vco_qtd)
        if "nome do contrato ou unidade" in p_text_norm:
            p.text = p.text.replace("[Nome do contrato ou unidade]", nome_contrato)
    
    # Replace literal remnants from template
    replace_regex_everywhere(doc, r"CONTRATO EXEMPLO - UNIDADE \(SETOR\)", nome_contrato)
    replace_regex_everywhere(doc, r"\(CR 00000 - 00000 - 00000 - CONTRATO EXEMPLO\)", nome_contrato)
    replace_regex_everywhere(doc, r"CONTRATO EXEMPLO", nome_contrato)
    replace_regex_everywhere(doc, r"13/07/2026 a 13/07/2026", periodo_str)
    
    if periodo_str:
        replace_regex_everywhere(doc, r"\[dd/mm/aaaa a dd/mm/aaaa\]", periodo_str)
        replace_regex_everywhere(doc, r"no período de \[dd/mm/aaaa a dd/mm/aaaa\]", f"no período de {periodo_str}")
        replace_regex_everywhere(doc, r"no período de ,", "")
    else:
        replace_regex_everywhere(doc, r"\[dd/mm/aaaa a dd/mm/aaaa\]", "todo o período")
        replace_regex_everywhere(doc, r"no período de \[dd/mm/aaaa a dd/mm/aaaa\],", "")
        replace_regex_everywhere(doc, r"no período de ,", "")
        
    replace_regex_everywhere(doc, r"\[Mês/Ano\s*[–-]\s*Mês/Ano\]", periodo_mes_ano)
    replace_regex_everywhere(doc, r"\[Nome completo\s*[–-]\s*registro profissional\]", resp_tecnico)
    replace_regex_everywhere(doc, r"\[Nome completo\]", resp_tecnico)
    replace_regex_everywhere(doc, r"\[Registro profissional\]", "CRM/COREN/CRP/CRP")
    replace_regex_everywhere(doc, r"\[Cargo\s*[–-]\s*Ex\.:\s*Gerente\s+SSMA\s*/\s*Médico\s+Coordenador\s+PCMSO\]", "Responsável Técnico SSMA")
    replace_regex_everywhere(doc, r"\[Cargo\s*/\s*Função\]", "\nResponsável Técnico SSMA")

    # 22. Standardize non-heading fonts to match template default font
    try:
        format_non_heading_text(doc)
    except Exception as e:
        print(f"Error standardizing fonts: {e}")

    # Save output file
    output_dir = os.path.join(settings.MEDIA_ROOT, "psicossocial", "resultados")
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"relatorio_psicossocial_{projeto.id}_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    
    return f"psicossocial/resultados/{filename}"


